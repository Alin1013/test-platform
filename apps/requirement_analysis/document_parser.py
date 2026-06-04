from dataclasses import dataclass, field
from io import BytesIO
import json
from pathlib import Path
from typing import Callable, Optional
import zipfile
from xml.etree import ElementTree

from bs4 import BeautifulSoup
from PIL import Image

try:
    from PyPDF2 import PdfReader
except ImportError:  # pragma: no cover - compatibility for older PyPDF2
    from PyPDF2 import PdfFileReader as PdfReader

try:
    import docx
except ImportError:  # pragma: no cover - dependency is declared in requirements
    docx = None


VisionExtractor = Callable[[str, str, bytes], str]


@dataclass
class ParsedDocument:
    text: str
    file_type: str
    metadata: dict = field(default_factory=dict)


class UnsupportedSourceFileError(ValueError):
    pass


class EmptyExtractedTextError(ValueError):
    pass


class VisionExtractorRequiredError(ValueError):
    pass


class DocumentParser:
    SUPPORTED_EXTENSIONS = {
        ".txt",
        ".md",
        ".pdf",
        ".docx",
        ".html",
        ".htm",
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".xmind",
    }
    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp"}

    @classmethod
    def extract_from_bytes(
        cls,
        filename: str,
        data: bytes,
        vision_extractor: Optional[VisionExtractor] = None,
    ) -> ParsedDocument:
        extension = Path(filename).suffix.lower()
        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise UnsupportedSourceFileError(f"不支持的源文件类型: {extension or filename}")

        if extension in {".txt", ".md"}:
            text = cls._extract_text(data)
            file_type = extension.lstrip(".")
            metadata = {"encoding": cls._detect_encoding(data)}
        elif extension in {".html", ".htm"}:
            text = cls._extract_html(data)
            file_type = "html"
            metadata = {}
        elif extension == ".pdf":
            text = cls._extract_pdf(data)
            file_type = "pdf"
            metadata = {}
        elif extension == ".docx":
            text = cls._extract_docx(data)
            file_type = "docx"
            metadata = {}
        elif extension == ".xmind":
            text, metadata = cls._extract_xmind(data)
            file_type = "xmind"
        else:
            text, metadata = cls._extract_image(filename, extension, data, vision_extractor)
            file_type = extension.lstrip(".")

        text = (text or "").strip()
        if not text:
            raise EmptyExtractedTextError("源文件未解析出有效文本内容")

        return ParsedDocument(text=text, file_type=file_type, metadata=metadata)

    @staticmethod
    def _detect_encoding(data: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gbk"):
            try:
                data.decode(encoding)
                return encoding
            except UnicodeDecodeError:
                continue
        return "utf-8"

    @classmethod
    def _extract_text(cls, data: bytes) -> str:
        encoding = cls._detect_encoding(data)
        return data.decode(encoding, errors="replace")

    @classmethod
    def _extract_html(cls, data: bytes) -> str:
        html = cls._extract_text(data)
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.extract()
        return soup.get_text(separator="\n", strip=True)

    @staticmethod
    def _extract_pdf(data: bytes) -> str:
        reader = PdfReader(BytesIO(data))
        parts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text.strip())
        return "\n".join(parts)

    @staticmethod
    def _extract_docx(data: bytes) -> str:
        if docx is None:
            raise UnsupportedSourceFileError("当前环境缺少 python-docx，无法解析 Word 文档")

        document = docx.Document(BytesIO(data))
        parts = []
        parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    @classmethod
    def _extract_xmind(cls, data: bytes) -> tuple[str, dict]:
        try:
            with zipfile.ZipFile(BytesIO(data)) as archive:
                names = set(archive.namelist())
                if "content.json" in names:
                    payload = json.loads(archive.read("content.json").decode("utf-8"))
                    points = cls._extract_xmind_json_points(payload)
                elif "content.xml" in names:
                    root = ElementTree.fromstring(archive.read("content.xml"))
                    points = cls._extract_xmind_xml_points(root)
                else:
                    raise EmptyExtractedTextError("XMind 文件缺少 content.json 或 content.xml")
        except zipfile.BadZipFile as exc:
            raise UnsupportedSourceFileError("XMind 文件格式无效") from exc
        except json.JSONDecodeError as exc:
            raise UnsupportedSourceFileError("XMind content.json 格式无效") from exc
        except ElementTree.ParseError as exc:
            raise UnsupportedSourceFileError("XMind content.xml 格式无效") from exc

        if not points:
            raise EmptyExtractedTextError("XMind 文件未解析出测试点")

        text = "\n".join(point["source_trace"] for point in points)
        return text, {"xmind_test_points": points}

    @classmethod
    def _extract_xmind_json_points(cls, payload) -> list[dict]:
        if isinstance(payload, dict):
            sheets = payload.get("sheets") or payload.get("children") or [payload]
        elif isinstance(payload, list):
            sheets = payload
        else:
            sheets = []

        raw_points = []
        for sheet in sheets:
            if not isinstance(sheet, dict):
                continue
            root_topic = sheet.get("rootTopic") or sheet.get("root_topic") or sheet.get("topic")
            if root_topic:
                cls._collect_xmind_json_leaves(root_topic, [], raw_points)

        return cls._format_xmind_points(raw_points)

    @classmethod
    def _collect_xmind_json_leaves(cls, topic: dict, parents: list[str], raw_points: list[dict]):
        title = str(topic.get("title") or "").strip()
        path = parents + ([title] if title else [])
        children = cls._xmind_json_children(topic)
        if children:
            for child in children:
                cls._collect_xmind_json_leaves(child, path, raw_points)
            return

        if title:
            raw_points.append({
                "title": title,
                "source_trace": " > ".join(path),
                "test_object": parents[-1] if parents else title,
                "note": cls._xmind_json_note(topic),
            })

    @staticmethod
    def _xmind_json_children(topic: dict) -> list[dict]:
        children = topic.get("children") or {}
        result = []
        if isinstance(children, dict):
            for value in children.values():
                if isinstance(value, list):
                    result.extend(item for item in value if isinstance(item, dict))
                elif isinstance(value, dict):
                    result.append(value)
        elif isinstance(children, list):
            result.extend(item for item in children if isinstance(item, dict))
        return result

    @staticmethod
    def _xmind_json_note(topic: dict) -> str:
        notes = topic.get("notes")
        if isinstance(notes, dict):
            plain = notes.get("plain")
            if isinstance(plain, dict):
                return str(plain.get("content") or "").strip()
            if isinstance(plain, str):
                return plain.strip()
        return ""

    @classmethod
    def _extract_xmind_xml_points(cls, root) -> list[dict]:
        raw_points = []
        for sheet in cls._xml_children(root, "sheet"):
            for topic in cls._xml_children(sheet, "topic"):
                cls._collect_xmind_xml_leaves(topic, [], raw_points)
        if not raw_points and cls._xml_name(root.tag) == "topic":
            cls._collect_xmind_xml_leaves(root, [], raw_points)
        return cls._format_xmind_points(raw_points)

    @classmethod
    def _collect_xmind_xml_leaves(cls, topic, parents: list[str], raw_points: list[dict]):
        title = cls._xml_child_text(topic, "title")
        path = parents + ([title] if title else [])
        children = cls._xmind_xml_topic_children(topic)
        if children:
            for child in children:
                cls._collect_xmind_xml_leaves(child, path, raw_points)
            return

        if title:
            raw_points.append({
                "title": title,
                "source_trace": " > ".join(path),
                "test_object": parents[-1] if parents else title,
                "note": cls._xmind_xml_note(topic),
            })

    @classmethod
    def _xmind_xml_topic_children(cls, topic) -> list:
        children = []
        for children_node in cls._xml_children(topic, "children"):
            for topics_node in cls._xml_children(children_node, "topics"):
                children.extend(cls._xml_children(topics_node, "topic"))
        return children

    @classmethod
    def _xmind_xml_note(cls, topic) -> str:
        for notes_node in cls._xml_children(topic, "notes"):
            for plain_node in cls._xml_children(notes_node, "plain"):
                return "".join(plain_node.itertext()).strip()
        return ""

    @classmethod
    def _format_xmind_points(cls, raw_points: list[dict]) -> list[dict]:
        points = []
        for index, point in enumerate(raw_points, 1):
            note = point.get("note", "")
            points.append({
                "id": f"TP-XMIND-{index:03d}",
                "title": point["title"],
                "test_object": point.get("test_object") or point["title"],
                "coverage_type": "功能",
                "design_technique": "XMind导入",
                "priority": "P2",
                "preconditions": "",
                "test_data_hint": note,
                "expected_focus": note or point["title"],
                "source_trace": point["source_trace"],
                "review_status": "pending",
                "review_comment": "",
            })
        return points

    @staticmethod
    def _xml_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    @classmethod
    def _xml_children(cls, node, name: str) -> list:
        return [child for child in list(node) if cls._xml_name(child.tag) == name]

    @classmethod
    def _xml_child_text(cls, node, name: str) -> str:
        for child in cls._xml_children(node, name):
            return "".join(child.itertext()).strip()
        return ""

    @staticmethod
    def _extract_image(
        filename: str,
        extension: str,
        data: bytes,
        vision_extractor: Optional[VisionExtractor],
    ) -> tuple[str, dict]:
        if vision_extractor is None:
            raise VisionExtractorRequiredError("图片源文件需要配置视觉解析模型")

        content_type = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".bmp": "image/png",
        }[extension]
        image_data = data
        metadata = {"vision_content_type": content_type}

        if extension == ".bmp":
            with Image.open(BytesIO(data)) as image:
                converted = BytesIO()
                image.convert("RGB").save(converted, format="PNG")
                image_data = converted.getvalue()
            metadata["converted_from"] = "bmp"

        return vision_extractor(filename, content_type, image_data), metadata
